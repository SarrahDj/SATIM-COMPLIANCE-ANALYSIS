import os
import re
import json
import csv
import numpy as np
import faiss
import pickle
from typing import Dict, List, Union, Optional
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
import spacy
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd

# Ensure spacy model is downloaded
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Downloading spaCy model 'en_core_web_sm'...")
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

class ComplianceResult:
    """
    Represents the result of a compliance check for a given use case.
    """
    def __init__(self, compliance_status: str, kpi_score: float, 
                 policy_matches: List[Dict], regulation_matches: List[Dict]):
        self.compliance_status = compliance_status
        self.kpi_score = kpi_score
        self.policy_matches = policy_matches
        self.regulation_matches = regulation_matches

class ComplianceProcessor:
    """
    Processes compliance documents (policies and regulations) and evaluates use cases
    against them to determine compliance status and KPI scores.
    """
    def __init__(self, model_name: str = 'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2'):
        self.model = SentenceTransformer(model_name)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Increased for better context
            chunk_overlap=100,
            separators=["\n\n", "\n", "shall:", "must:", "should:", "requirements:", ". "]
        )
        self.nlp = nlp
        self.policy_index = None
        self.regulation_index = None
        self.policy_metadata = []
        self.regulation_metadata = []
        self.dimension = self.model.get_sentence_embedding_dimension()
        
        # Enhanced domain-specific keywords for better classification
        self.domain_keywords = {
            'classification': ['classification', 'asset', 'inventory', 'ownership', 'labeling', 'sensitivity', 'data handling'],
            'logging': ['log', 'audit', 'monitoring', 'retention', 'review', 'activity', 'event recording', 'system events'],
            'access_review': ['access', 'review', 'permissions', 'privileges', 'authorization', 'entitlements', 'user management', 'identity'],
            'backup': ['backup', 'recovery', 'restore', 'disaster', 'continuity', 'archiving', 'data resilience', 'business continuity'],
            'endpoint': ['endpoint', 'device', 'workstation', 'antivirus', 'firewall', 'mobile', 'security controls'],
            'vulnerability': ['vulnerability', 'patch', 'scan', 'remediation', 'cve', 'security update', 'threat management', 'risk assessment'],
            'training': ['training', 'awareness', 'education', 'employee', 'security awareness', 'compliance training', 'personnel development']
        }
        
        # Standard to domain mapping for specific regulatory standards
        self.standard_mapping = {
            "ISO 27001 - A.8.2": "classification",
            "PCI DSS Req. 10": "logging",
            "NIST AU-2": "logging",
            "ISO 27001 - A.9.2.5": "access_review",
            "ISO 27001 - A.12.3": "backup",
            "NIST 800-53 - CM-7": "endpoint",
            "ISO 27001 - A.12.6.1": "vulnerability",
            "ISO 27001 - A.7.2.2": "training"
        }

    def load_document(self, file_path: str) -> str:
        """
        Loads text content from various document types (PDF, DOCX, TXT).
        """
        try:
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_path}")
        except Exception as e:
            print(f"Error loading document '{file_path}': {str(e)}")
            return ""

    def preprocess_text(self, text: str, doc_type: str) -> str:
        """
        Cleans and preprocesses text, removing boilerplate and adding domain-specific markers.
        """
        # Remove common headers/footers
        text = re.sub(r'Policy #:.+?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Title:.+?\n', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\d+\.\d+(\.\d+)*\s*', '', text)  # Remove numeric outlines
        text = re.sub(r'\[[\w\s\-,]+\]', '', text)  # Remove bracketed references
        text = re.sub(r'\(\w+\s+\d{4}\)', '', text)  # Remove (Word Year) patterns
        text = re.sub(r'\s+', ' ', text) # Replace multiple spaces with single space
        
        # Domain-specific enhancements for policy and regulation documents
        if doc_type == "policy":
            text = re.sub(r'\b(shall|must|is required to)\b', r'REQUIREMENT: \1', text, flags=re.IGNORECASE)
        elif doc_type == "regulation":
            text = re.sub(r'\b(compliance|shall not|violation|penalty)\b', r'REGULATION: \1', text, flags=re.IGNORECASE)
            
        return text.strip()

    def process_documents(self, policy_dir: str, regulation_dir: str) -> bool:
        """
        Processes documents from specified directories, chunks them, generates embeddings,
        and builds FAISS indexes for efficient similarity search.
        """
        # Process policies
        if os.path.exists(policy_dir):
            policy_chunks = []
            policy_metadata = []
            
            print(f"Processing policy documents from: {policy_dir}")
            for root, _, files in os.walk(policy_dir):
                for file in files:
                    if file.split('.')[-1].lower() not in ['pdf', 'docx', 'txt']:
                        print(f"Skipping unsupported file: {file}")
                        continue
                        
                    path = os.path.join(root, file)
                    text = self.load_document(path)
                    if not text:
                        continue # Skip empty documents
                        
                    processed_text = self.preprocess_text(text, "policy")
                    doc_chunks = self.text_splitter.split_text(processed_text)
                    
                    for i, chunk in enumerate(doc_chunks):
                        if len(chunk.strip()) < 50: # Skip very short chunks
                            continue
                            
                        # Determine domain based on keywords in the chunk
                        domain = "general"
                        chunk_lower = chunk.lower()
                        for dom, keywords in self.domain_keywords.items():
                            if any(kw in chunk_lower for kw in keywords):
                                domain = dom
                                break
                        
                        policy_metadata.append({
                            "source": file,
                            "text": chunk,
                            "domain": domain,
                            "type": "policy"
                        })
                        policy_chunks.append(chunk)
            
            if policy_chunks:
                print(f"Generating embeddings for {len(policy_chunks)} policy chunks...")
                policy_embeddings = self.model.encode(policy_chunks, show_progress_bar=True)
                policy_embeddings = policy_embeddings.astype(np.float32)
                faiss.normalize_L2(policy_embeddings) # Normalize for inner product search
                
                self.policy_index = faiss.IndexFlatIP(self.dimension)
                self.policy_index.add(policy_embeddings)
                self.policy_metadata = policy_metadata
                print("Policy index created successfully!")
            else:
                print("No valid policy chunks found to create an index.")
        else:
            print(f"Policy directory not found: {policy_dir}")

        # Process regulations
        if os.path.exists(regulation_dir):
            regulation_chunks = []
            regulation_metadata = []
            
            print(f"Processing regulation documents from: {regulation_dir}")
            for root, _, files in os.walk(regulation_dir):
                for file in files:
                    if file.split('.')[-1].lower() not in ['pdf', 'docx', 'txt']:
                        print(f"Skipping unsupported file: {file}")
                        continue
                        
                    path = os.path.join(root, file)
                    text = self.load_document(path)
                    if not text:
                        continue # Skip empty documents
                        
                    processed_text = self.preprocess_text(text, "regulation")
                    doc_chunks = self.text_splitter.split_text(processed_text)
                    
                    for i, chunk in enumerate(doc_chunks):
                        if len(chunk.strip()) < 50: # Skip very short chunks
                            continue
                            
                        # Determine domain based on keywords in the chunk
                        domain = "general"
                        chunk_lower = chunk.lower()
                        for dom, keywords in self.domain_keywords.items():
                            if any(kw in chunk_lower for kw in keywords):
                                domain = dom
                                break
                        
                        regulation_metadata.append({
                            "source": file,
                            "text": chunk,
                            "domain": domain,
                            "type": "regulation"
                        })
                        regulation_chunks.append(chunk)
            
            if regulation_chunks:
                print(f"Generating embeddings for {len(regulation_chunks)} regulation chunks...")
                regulation_embeddings = self.model.encode(regulation_chunks, show_progress_bar=True)
                regulation_embeddings = regulation_embeddings.astype(np.float32)
                faiss.normalize_L2(regulation_embeddings) # Normalize for inner product search
                
                self.regulation_index = faiss.IndexFlatIP(self.dimension)
                self.regulation_index.add(regulation_embeddings)
                self.regulation_metadata = regulation_metadata
                print("Regulation index created successfully!")
            else:
                print("No valid regulation chunks found to create an index.")
        else:
            print(f"Regulation directory not found: {regulation_dir}")
        
        return self.policy_index is not None or self.regulation_index is not None

    def match_usecase(self, usecase: str, standard_ref: str = "", top_k: int = 5) -> ComplianceResult:
        """
        Matches a given use case against indexed policies and regulations, calculates KPI,
        and determines compliance status.
        """
        policy_matches = []
        regulation_matches = []
        
        # Determine target domain from standard reference if provided
        target_domain = self.standard_mapping.get(standard_ref, "")
        
        # Create an enhanced query by extracting keywords and adding domain-specific terms
        doc = self.nlp(usecase)
        keywords = [token.lemma_ for token in doc if token.pos_ in ['NOUN', 'ADJ', 'VERB'] and not token.is_stop]
        
        if target_domain:
            keywords.extend(self.domain_keywords[target_domain])
        
        query_text = f"COMPLIANCE USE CASE: {usecase} KEYWORDS: {' '.join(set(keywords))}"
        if standard_ref:
            query_text += f" STANDARD: {standard_ref}"
            
        query_embedding = self.model.encode([query_text])
        query_embedding = query_embedding.astype(np.float32)
        faiss.normalize_L2(query_embedding)
        
        # Search policy index for relevant matches
        if self.policy_index:
            distances, indices = self.policy_index.search(query_embedding, top_k)
            for i in range(top_k):
                idx = indices[0][i]
                if idx < 0 or idx >= len(self.policy_metadata): # Ensure index is valid
                    continue
                    
                meta = self.policy_metadata[idx]
                policy_matches.append({
                    "source": meta["source"],
                    "text": meta["text"][:200] + "...", # Truncate text for display
                    "score": float(distances[0][i]),
                    "domain": meta["domain"]
                })
        
        # Search regulation index for relevant matches
        if self.regulation_index:
            distances, indices = self.regulation_index.search(query_embedding, top_k)
            for i in range(top_k):
                idx = indices[0][i]
                if idx < 0 or idx >= len(self.regulation_metadata): # Ensure index is valid
                    continue
                    
                meta = self.regulation_metadata[idx]
                regulation_matches.append({
                    "source": meta["source"],
                    "text": meta["text"][:200] + "...", # Truncate text for display
                    "score": float(distances[0][i]),
                    "domain": meta["domain"]
                })
        
        # Calculate KPI based on matched documents
        all_matches = policy_matches + regulation_matches
        kpi_score = 0.0
        
        if all_matches:
            # Calculate weighted average score, giving higher priority to top matches
            scores = sorted([match["score"] for match in all_matches], reverse=True)
            weights = [0.4, 0.3, 0.2, 0.1] # Weights for top 4 matches
            
            weighted_sum = 0
            total_weight = 0
            for i, score in enumerate(scores[:len(weights)]):
                weighted_sum += score * weights[i]
                total_weight += weights[i]
                
            avg_score = weighted_sum / total_weight if total_weight > 0 else (scores[0] if scores else 0)
            
            # Convert similarity score to KPI (0-100 scale)
            kpi_score = min(100, max(0, avg_score * 100)) # Scale to 0-100 and clip
            
            # Boost score if the target domain has strong matches
            if target_domain:
                domain_matches = [m for m in all_matches if m["domain"] == target_domain]
                if domain_matches:
                    kpi_score = min(100, kpi_score + 5) # Small boost for domain relevance
            
            # Penalize if no regulation matches for domains where regulations are critical
            regulation_heavy = ["logging", "access_review", "vulnerability"]
            if target_domain in regulation_heavy and not regulation_matches:
                kpi_score = max(0, kpi_score - 15) # Significant penalty for missing regulations
        
        # Determine compliance status based on KPI score thresholds
        if kpi_score >= 80:
            status = "COMPLIANT"
        elif kpi_score >= 65:
            status = "PARTIAL_COMPLIANT"
        else:
            status = "NON_COMPLIANT"
            
        return ComplianceResult(status, kpi_score, policy_matches, regulation_matches)

def evaluate_use_cases(docx_path: str, output_csv: str, policy_dir: str, regulation_dir: str):
    """
    Main function to evaluate compliance use cases and save results to CSV.
    """
    processor = ComplianceProcessor()
    # Process documents and build indexes
    if not processor.process_documents(policy_dir, regulation_dir):
        print("Failed to process documents! Ensure policy and regulation directories exist and contain valid files.")
        return

    # Load use cases from DOCX table
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error loading DOCX document '{docx_path}': {str(e)}")
        return

    if not doc.tables:
        print("No tables found in the DOCX document. Please ensure your use cases are in a table.")
        return

    table = doc.tables[0]
    results = [] # List to store dictionaries of results
    
    # Mapping from French to English compliance status
    status_map = {
        "Conforme": "COMPLIANT",
        "Non conforme": "NON_COMPLIANT",
        "Partiellement conforme": "PARTIAL_COMPLIANT"
    }

    # Process each use case, skipping the header row (index 0)
    print("Processing use cases from DOCX table...")
    for row_idx, row in enumerate(table.rows[1:]):
        cells = row.cells
        if len(cells) < 4:
            print(f"Skipping row {row_idx + 2} due to insufficient columns.") # +2 for header row and 0-indexing
            continue
            
        use_case = cells[0].text.strip()
        standard_ref = cells[1].text.strip()
        
        try:
            actual_kpi = float(cells[2].text.strip().replace('%', ''))
        except ValueError:
            print(f"Warning: Could not parse Actual KPI in row {row_idx + 2} ('{cells[2].text.strip()}'). Setting to 0.")
            actual_kpi = 0.0

        actual_status_raw = cells[3].text.strip()
        actual_status = status_map.get(actual_status_raw, "UNKNOWN")
        if actual_status == "UNKNOWN":
            print(f"Warning: Unknown Actual Status in row {row_idx + 2} ('{actual_status_raw}').")
        
        try:
            result = processor.match_usecase(use_case, standard_ref)
            results.append({
                "use_case": use_case,
                "standard_ref": standard_ref,
                "predicted_kpi": result.kpi_score,
                "predicted_status": result.compliance_status,
                "actual_kpi": actual_kpi,
                "actual_status": actual_status,
                "policy_matches": len(result.policy_matches),
                "regulation_matches": len(result.regulation_matches)
            })
        except Exception as e:
            print(f"Error processing use case '{use_case}' in row {row_idx + 2}: {str(e)}")
            results.append({
                "use_case": use_case,
                "standard_ref": standard_ref,
                "predicted_kpi": 0,
                "predicted_status": "ERROR",
                "actual_kpi": actual_kpi,
                "actual_status": actual_status,
                "policy_matches": 0,
                "regulation_matches": 0
            })

    # Write results to CSV
    if results:
        with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['use_case', 'standard_ref', 'predicted_kpi', 'predicted_status', 
                          'actual_kpi', 'actual_status', 'policy_matches', 'regulation_matches']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(results)
        print(f"Compliance results saved to '{output_csv}'")
        
        # Calculate and print accuracy metrics
        correct_status = 0
        kpi_error = 0
        total_status_evaluable = 0
        total_kpi_evaluable = 0
        
        for result in results:
            if result['actual_status'] != "UNKNOWN" and result['predicted_status'] != "ERROR":
                total_status_evaluable += 1
                if result['predicted_status'] == result['actual_status']:
                    correct_status += 1
            if result['predicted_status'] != "ERROR":
                total_kpi_evaluable += 1
                kpi_error += abs(result['predicted_kpi'] - result['actual_kpi'])
        
        if total_status_evaluable > 0:
            status_accuracy = correct_status / total_status_evaluable
            print(f"\nEvaluation Results:")
            print(f"Status Accuracy: {status_accuracy:.1%}")
        else:
            print("\nNo valid use cases for status accuracy calculation.")

        if total_kpi_evaluable > 0:
            avg_kpi_error = kpi_error / total_kpi_evaluable
            print(f"Average KPI Error: {avg_kpi_error:.1f} percentage points")
        else:
            print("No valid use cases for KPI error calculation.")
            
        # Show improvement tips
        print("\nRecommendations for improvement:")
        print("- Add more policy documents covering all compliance domains.")
        print("- Ensure regulations are included for standards like PCI DSS and NIST.")
        print("- Review document quality for key terms: 'shall', 'must', 'requirement'.")
        print("- Consider refining domain keywords for better use case classification.")
    else:
        print("\nNo valid use cases processed. No CSV generated.")
    

def generate_graphs_from_csv(csv_path: str, output_dir: str = "graphs"):
    """
    Generates various compliance-related graphs from a CSV file
    and saves them to the specified directory.

    Args:
        csv_path (str): The path to the CSV file containing compliance results.
        output_dir (str): The directory where generated graphs will be saved.
    """
    try:
        results_df = pd.read_csv(csv_path)
        if results_df.empty:
            print(f"The CSV file '{csv_path}' is empty or could not be read. No graphs generated.")
            return
    except FileNotFoundError:
        print(f"Error: CSV file not found at '{csv_path}'. Please ensure the compliance analyzer has been run first.")
        return
    except Exception as e:
        print(f"Error reading CSV file '{csv_path}': {str(e)}")
        return

    os.makedirs(output_dir, exist_ok=True) # Create output directory if it doesn't exist

    # Set a professional style for plots
    sns.set_theme(style="whitegrid", palette="pastel")
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.sans-serif'] = ['Inter', 'Arial']
    plt.rcParams['axes.labelsize'] = 12
    plt.rcParams['xtick.labelsize'] = 10
    plt.rcParams['ytick.labelsize'] = 10
    plt.rcParams['figure.titlesize'] = 14
    plt.rcParams['figure.autolayout'] = True # Automatically adjust subplot params for tight layout

    # 1. Predicted vs. Actual KPI Score by Use Case (Bar Chart)
    plt.figure(figsize=(12, 7))
    x = np.arange(len(results_df['use_case']))
    width = 0.35 # Width of the bars

    plt.bar(x - width/2, results_df['predicted_kpi'], width, label='Predicted KPI', color='#61A4E7')
    plt.bar(x + width/2, results_df['actual_kpi'], width, label='Actual KPI', color='#ADD8E6', alpha=0.8) # Lighter shade for actual

    plt.axhline(y=95, color='forestgreen', linestyle='--', label='Target KPI (95%)') # Target line
    plt.ylabel('KPI Score (%)')
    plt.title('Compliance KPI by Use Case: Predicted vs. Actual')
    plt.xticks(x, results_df['use_case'], rotation=45, ha='right')
    plt.ylim(0, 105) # Extend y-axis slightly above 100
    plt.legend()
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'kpi_by_use_case.png'))
    plt.close()

    # 2. Compliance Status Distribution (Pie Chart)
    plt.figure(figsize=(8, 8))
    status_counts = results_df['predicted_status'].value_counts()
    colors = {'COMPLIANT': '#8FD085', 'PARTIAL_COMPLIANT': '#FFD700', 'NON_COMPLIANT': '#DC143C', 'ERROR': '#808080', 'UNKNOWN': '#C0C0C0'} # Green, Yellow, Red, Gray for errors/unknowns
    
    # Ensure all possible statuses are covered for consistent coloring
    all_statuses_possible = ["COMPLIANT", "PARTIAL_COMPLIANT", "NON_COMPLIANT", "ERROR", "UNKNOWN"]
    # Filter status_counts to include only present statuses, then reindex to maintain order and add missing ones with 0
    plot_data = status_counts.reindex(all_statuses_possible, fill_value=0)
    # Filter out statuses that have zero count for cleaner pie chart
    plot_data = plot_data[plot_data > 0] 

    # If plot_data is empty after filtering, don't attempt to draw pie chart
    if not plot_data.empty:
        wedges, texts, autotexts = plt.pie(
            plot_data,
            labels=plot_data.index,
            autopct='%1.1f%%',
            startangle=90,
            colors=[colors.get(s, '#CCCCCC') for s in plot_data.index], # Use get with a default color for safety
            pctdistance=0.85, # Distance of the percentage labels
            explode=[0.05 if s == 'NON_COMPLIANT' else 0 for s in plot_data.index] # Slightly explode 'NON_COMPLIANT'
        )
        plt.title('Distribution of Predicted Compliance Statuses')
        plt.setp(autotexts, size=10, weight="bold", color="white") # Make percentages bold and white
        plt.setp(texts, size=11, color="dimgray") # Make labels a bit larger
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'compliance_status_distribution.png'))
    else:
        print("No valid compliance statuses to display in the pie chart.")
    plt.close()

    # 3. Policy and Regulation Matches by Use Case (Stacked Bar Chart)
    plt.figure(figsize=(12, 7))
    
    # Ensure columns exist before plotting
    if 'policy_matches' in results_df.columns and 'regulation_matches' in results_df.columns:
        results_df[['policy_matches', 'regulation_matches']].plot(
            kind='bar', stacked=True, ax=plt.gca(), # Use current axes
            color=['#7B68EE', '#B0C4DE'] # MediumSlateBlue, LightSteelBlue
        )
        plt.ylabel('Number of Matches')
        plt.title('Policy and Regulation Matches per Use Case')
        plt.xticks(ticks=range(len(results_df)), labels=results_df['use_case'], rotation=45, ha='right')
        plt.legend(['Policy Matches', 'Regulation Matches'], title='Document Type')
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'matches_by_use_case.png'))
    else:
        print("Columns 'policy_matches' or 'regulation_matches' not found in CSV. Cannot generate stacked bar chart.")
    plt.close()

    print(f"Graphs saved to '{output_dir}' directory.")


if __name__ == "__main__":
    # Configure paths
    # IMPORTANT: Ensure 'policies' and 'regulations' directories exist in the same
    # location as this script, and contain your .pdf, .docx, or .txt files.
    # Replace "Evaluation_GRC_Use_Cases_1_8(2) (1).docx" with the actual path to your DOCX file.
    POLICY_DIR = "policies" 
    REGULATION_DIR = "regulations"
    DOCX_PATH = "Evaluation_GRC_Use_Cases_1_8(2) (1).docx" 
    OUTPUT_CSV = "compliance_results.csv"

    # --- IMPORTANT: This section previously generated dummy files. ---
    # --- It has been removed. Please ensure these files/directories ---
    # --- exist before running the script. ---
    # Example:
    # policies/
    #   security_policy.txt
    # regulations/
    #   gdpr_overview.txt
    # Evaluation_GRC_Use_Cases_1_8(2) (1).docx
    
    # Run evaluation
    evaluate_use_cases(DOCX_PATH, OUTPUT_CSV, POLICY_DIR, REGULATION_DIR)
    
    # Define the path to the CSV file generated by your compliance analyzer script
    CSV_FILE_PATH = "compliance_results.csv"
    # Define the directory where the graphs will be saved
    GRAPH_OUTPUT_DIR = "graphs"

    generate_graphs_from_csv(CSV_FILE_PATH, GRAPH_OUTPUT_DIR)
